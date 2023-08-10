from docx import Document
from io import BytesIO

TemplateValue = dict[str, str]


class DocumentsFormatter:
    def __init__(self, path, templates_values: dict):
        """
        Format document and return
        :param templates_values is dict same `{'template':str, 'value': str}`
        """

        self.document = Document(path)
        self.templates_values = templates_values

    def format(self) -> BytesIO:
        self.__format_paragraphs()
        self.__format_tables()

        # Save the .docx to the buffer
        file_stream = BytesIO()
        self.document.save(file_stream)

        file_stream.seek(0)

        return file_stream

    def __format_tables(self):
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.__format_paragraphs(cell)

    def __format_paragraphs(self, obj=None):
        if obj is None:
            obj = self.document

        for p in obj.paragraphs:
            for template, value in self.templates_values.items():
                self.__format(p, template, value)

    def __format(self, string: str, template: str, value: str):
        string.text = string.text.replace(template, value)
