from docx import Document
from io import BytesIO


# thanks
# https://newbedev.com/python-docx-replace-string-in-paragraph-while-keeping-style

from typing import Dict, List


class KeyChanger:
    def __init__(self, p, key, value) -> None:
        self.p = p
        self.key = key
        self.value = value
        self.run_text = ""
        self.runs_indexes: List = []
        self.run_char_indexes: List = []
        self.runs_to_change: Dict = {}

    def _initialize(self) -> None:
        run_index = 0
        for run in self.p.runs:
            self.run_text += run.text
            self.runs_indexes += [run_index for _ in run.text]
            self.run_char_indexes += [char_index for char_index, char in enumerate(run.text)]
            run_index += 1

    def replace(self) -> None:
        self._initialize()
        parsed_key_length = len(self.key)
        index_to_replace = self.run_text.find(self.key)

        for i in range(parsed_key_length):
            index = index_to_replace + i
            run_index = self.runs_indexes[index]
            run = self.p.runs[run_index]
            run_char_index = self.run_char_indexes[index]

            if not self.runs_to_change.get(run_index):
                self.runs_to_change[run_index] = [char for char_index, char in enumerate(run.text)]

            run_to_change: Dict = self.runs_to_change.get(run_index)  # type: ignore[assignment]
            if index == index_to_replace:
                run_to_change[run_char_index] = self.value
            else:
                run_to_change[run_char_index] = ""

        # make the real replace
        for index, text in self.runs_to_change.items():
            run = self.p.runs[index]
            run.text = "".join(text)


class DocumentsFormatter:
    def __init__(self, path, templates_values):
        self.document = Document(path)
        self.data = self.__get_primitive_templates_values(
            templates_values
        )

    def format(self) -> BytesIO:
        print(self.data)
        for p in self.all_paragraphs:
            for key, val in self.data.items():
                while key in p.text:
                    KeyChanger(p, key, val).replace()

                # if key in p.text:
                #     print('\n' * 10, flush=True)
                #     print('Start find new key -', key, flush=True)
                #     inline = p.runs
                #     # Replace strings and retain the same style.
                #     # The text to be replaced can be split over several runs so
                #     # search through, identify which runs need to have text replaced
                #     # then replace the text in those identified
                #     started = False
                #     key_index = 0
                #     # found_runs is a list of (inline index, index of match, length of match)
                #     found_runs = list()
                #     found_all = False
                #     replace_done = False
                #     for i in range(len(inline)):
                #         print('start new inline', flush=True)
                #         # case 1: found in single run so short circuit the replace
                #         if key in inline[i].text and not started:
                #             found_runs.append((i, inline[i].text.find(key), len(key)))
                #             text = inline[i].text.replace(key, str(val))
                #             inline[i].text = text
                #             replace_done = True
                #             found_all = True
                #             break

                #         if key[key_index] not in inline[i].text and not started:
                #             # keep looking ...
                #             continue

                #         # case 2: search for partial text, find first run
                #         if key[key_index] in inline[i].text and inline[i].text[-1] in key and not started:
                #             # check sequence
                #             start_index = inline[i].text.find(key[key_index])
                #             check_length = len(inline[i].text)
                #             for text_index in range(start_index, check_length):
                #                 if inline[i].text[text_index] != key[key_index]:
                #                     # no match so must be false positive
                #                     break

                #             if key_index == 0:
                #                 started = True
                #             chars_found = check_length - start_index
                #             key_index += chars_found
                #             found_runs.append((i, start_index, chars_found))

                #             if key_index != len(key):
                #                 continue
                #             else:
                #                 print('found all is true(67 line), break')
                #                 # found all chars in key
                #                 found_all = True
                #                 break

                #         # case 2: search for partial text, find subsequent run
                #         if key[key_index] in inline[i].text and started and not found_all:
                #             # check sequence
                #             chars_found = 0
                #             check_length = len(inline[i].text)
                #             for text_index in range(0, check_length):
                #                 if key_index != len(key) and inline[i].text[text_index] == key[key_index]:
                #                     key_index += 1
                #                     chars_found += 1
                #                     print('key index up ->', key_index, flush=True)
                #                 else:
                #                     break

                #             # no match so must be end
                #             found_runs.append((i, 0, chars_found))
                #             print('key index, len key', key_index, len(key), flush=True)
                #             if key_index == len(key):
                #                 found_all = True
                #                 print('found all is true', flush=True)
                #                 print('replace done is', replace_done, flush=True)
                #                 break

                #     if found_all and not replace_done:
                #         print('replace', flush=True)
                #         for i, item in enumerate(found_runs):
                #             index, start, length = [t for t in item]
                #             if i == 0:
                #                 text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                #                 inline[index].text = text
                #                 print('replace')
                #             else:
                #                 text = inline[index].text.replace(inline[index].text[start:start + length], '')
                #                 inline[index].text = text
                #     else:
                #         print('not replace', flush=True)

        file_stream = BytesIO()
        self.document.save(file_stream)

        file_stream.seek(0)

        return file_stream

    def __get_primitive_templates_values(self, templates_values):
        res = {}
        for tv in templates_values:
            res[tv.template.name_in_document] = tv.value
        return res

    @property
    def all_paragraphs(self):
        all_paragraphs = list(self.document.paragraphs)
        for t in self.document.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_paragraphs.append(paragraph)

        return all_paragraphs
